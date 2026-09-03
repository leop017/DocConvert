import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document as DocxDocument

from docconvert.converters import DocConverter, WordConverter


class TestWordConverterSheetsRejection(unittest.TestCase):
    """WordConverter must reject the sheets= kwarg explicitly."""

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.docx = os.path.join(self.tmp.name, "t.docx")
        DocxDocument().save(self.docx)
        self.out_dir = Path(self.tmp.name) / "out"
        self.out_dir.mkdir()

    def test_word_converter_rejects_sheets(self):
        c = WordConverter()
        with self.assertRaises(TypeError) as ctx:
            c.convert(self.docx, "html", self.out_dir, sheets=["S1"])
        self.assertIn("sheets", str(ctx.exception).lower())

    def test_word_converter_accepts_no_sheets(self):
        c = WordConverter()
        results, errors = c.convert(self.docx, "html", self.out_dir)
        self.assertEqual(len(errors), 0)
        self.assertEqual(len(results), 1)


class TestDocConverterSheetsRejection(unittest.TestCase):
    """DocConverter must reject the sheets= kwarg explicitly."""

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.doc = os.path.join(self.tmp.name, "t.doc")
        Path(self.doc).write_bytes(b"placeholder")
        self.out_dir = Path(self.tmp.name) / "out"
        self.out_dir.mkdir()

    def test_doc_converter_rejects_sheets(self):
        c = DocConverter()
        with self.assertRaises(TypeError) as ctx:
            c.convert(self.doc, "html", self.out_dir, sheets=["S1"])
        self.assertIn("sheets", str(ctx.exception).lower())


class TestDocConverterJsonFormat(unittest.TestCase):
    """``DocConverter`` must defer JSON serialization to ``JsonExporter``
    instead of pre-serializing with ``json.dumps`` and handing the
    exporter a string — otherwise the exporter's ``indent=2`` and
    ``ensure_ascii=False`` flags are silently bypassed.

    textract is mocked because the .doc pipeline depends on a system
    binary (antiword) that is not always available in test environments.
    """

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.doc = os.path.join(self.tmp.name, "t.doc")
        Path(self.doc).write_bytes(b"placeholder")
        self.out_dir = Path(self.tmp.name) / "out"
        self.out_dir.mkdir()

    @unittest.skipUnless(
        __import__("importlib").util.find_spec("textract") is not None,
        "textract not installed (Linux/macOS only; Windows skips)",
    )
    def test_json_output_is_pretty_printed(self):
        import json
        from unittest.mock import patch
        c = DocConverter()
        with patch("textract.process", return_value=b"hello world"):
            results, errors = c.convert(self.doc, "json", self.out_dir)
        self.assertEqual(errors, [])
        self.assertEqual(len(results), 1)
        content = Path(results[0][1]).read_text(encoding="utf-8")
        # indent=2 in JsonExporter must produce a newline-delimited file
        self.assertIn("\n", content)
        # And the result must still be valid JSON
        parsed = json.loads(content)
        self.assertIn("source", parsed)
        self.assertIn("content", parsed)


class TestDocConverterHtmlAndMd(unittest.TestCase):
    """DocConverter HTML/MD paths must work when textract is mocked."""

    _has_textract = __import__("importlib").util.find_spec("textract") is not None

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.doc = os.path.join(self.tmp.name, "t.doc")
        Path(self.doc).write_bytes(b"placeholder")
        self.out_dir = Path(self.tmp.name) / "out"
        self.out_dir.mkdir()

    @unittest.skipUnless(_has_textract, "textract not installed (Linux/macOS only; Windows skips)")
    @patch("textract.process", return_value=b"hello world")
    def test_html_output_contains_doctype(self, _m):
        c = DocConverter()
        results, errors = c.convert(self.doc, "html", self.out_dir)
        self.assertEqual(errors, [])
        self.assertEqual(len(results), 1)
        content = Path(results[0][1]).read_text(encoding="utf-8")
        self.assertIn("<!DOCTYPE html>", content)
        self.assertIn("hello world", content)

    @unittest.skipUnless(_has_textract, "textract not installed (Linux/macOS only; Windows skips)")
    @patch("textract.process", return_value=b"hello world")
    def test_md_output_has_source_comment(self, _m):
        c = DocConverter()
        results, errors = c.convert(self.doc, "md", self.out_dir)
        self.assertEqual(errors, [])
        self.assertEqual(len(results), 1)
        content = Path(results[0][1]).read_text(encoding="utf-8")
        self.assertIn("<!-- source:", content)
        self.assertIn("hello world", content)

    @unittest.skipUnless(_has_textract, "textract not installed (Linux/macOS only; Windows skips)")
    @patch("textract.process", return_value=b"hello world")
    def test_enhanced_md_wraps_in_pre(self, _m):
        c = DocConverter()
        results, errors = c.convert(self.doc, "md", self.out_dir, enhanced_md=True)
        self.assertEqual(errors, [])
        content = Path(results[0][1]).read_text(encoding="utf-8")
        # html_to_md converts <pre> to fenced code block
        self.assertIn("```", content)
        self.assertIn("hello world", content)

    @unittest.skipUnless(_has_textract, "textract not installed (Linux/macOS only; Windows skips)")
    @patch("textract.process", side_effect=Exception("boom"))
    def test_conversion_error_is_reported(self, _m):
        c = DocConverter()
        results, errors = c.convert(self.doc, "html", self.out_dir)
        self.assertEqual(results, [])
        self.assertEqual(len(errors), 1)
        self.assertEqual(errors[0][0], "t.doc")
        self.assertIn("boom", errors[0][1])

    @unittest.skipUnless(_has_textract, "textract not installed (Linux/macOS only; Windows skips)")
    @patch("textract.process", return_value=b"hello world")
    def test_unsupported_format_raises(self, _m):
        c = DocConverter()
        with self.assertRaises(ValueError) as ctx:
            c._convert_doc(self.doc, "xml", self.out_dir)
        self.assertIn("不支持的输出格式", str(ctx.exception))

    def test_cancel_check_returns_empty(self):
        c = DocConverter()
        c.cancel_check = lambda: True
        results, errors = c.convert(self.doc, "html", self.out_dir)
        self.assertEqual(results, [])
        self.assertEqual(errors, [])


class TestWordConverterCancelCheck(unittest.TestCase):
    """WordConverter must honor cancel_check at convert() entry."""

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.docx = os.path.join(self.tmp.name, "t.docx")
        DocxDocument().save(self.docx)
        self.out_dir = Path(self.tmp.name) / "out"
        self.out_dir.mkdir()

    def test_word_cancel_returns_empty_before_running(self):
        c = WordConverter()
        c.cancel_check = lambda: True
        results, errors = c.convert(self.docx, "html", self.out_dir)
        self.assertEqual(results, [])
        self.assertEqual(errors, [])


class TestWordConverterJsonFormat(unittest.TestCase):
    """``WordConverter._to_json`` must defer JSON serialization to
    ``JsonExporter`` (same contract as ``DocConverter``).
    """

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.docx = os.path.join(self.tmp.name, "t.docx")
        d = DocxDocument()
        d.add_paragraph("Hello world")
        d.save(self.docx)
        self.out_dir = Path(self.tmp.name) / "out"
        self.out_dir.mkdir()

    def test_to_json_returns_dict_not_string(self):
        """_to_json must return a dict so JsonExporter can serialize it."""
        import io
        c = WordConverter()
        d = DocxDocument(self.docx)
        buf = io.BytesIO()
        d.save(buf)
        buf.seek(0)
        result = c._to_json(buf, self.docx, 1)
        self.assertIsInstance(result, dict)
        self.assertIn("metadata", result)
        self.assertIn("content", result)

    def test_json_output_is_pretty_printed(self):
        """The exported JSON file must have indent=2 (pretty-printed)."""
        c = WordConverter()
        results, errors = c.convert(self.docx, "json", self.out_dir)
        self.assertEqual(errors, [])
        self.assertEqual(len(results), 1)
        content = Path(results[0][1]).read_text(encoding="utf-8")
        # indent=2 produces newline-delimited output
        self.assertIn("\n", content)
        import json
        parsed = json.loads(content)
        self.assertIn("metadata", parsed)
        self.assertIn("content", parsed)


class TestWordConverterTableHeaderDetection(unittest.TestCase):
    """WordConverter must respect Word's <w:tblHeader/> flag for tables.

    Regression: BUG 6 - first row of every table was force-converted
    to <th>, even when the source row was plain data.
    """

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.out_dir = Path(self.tmp.name) / "out"
        self.out_dir.mkdir()

    def _build_table_docx(self, name, mark_first_row_as_header: bool) -> str:
        from docx.oxml.ns import qn
        path = os.path.join(self.tmp.name, name)
        doc = DocxDocument()
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "h1" if mark_first_row_as_header else "row1col1"
        t.cell(0, 1).text = "h2" if mark_first_row_as_header else "row1col2"
        t.cell(1, 0).text = "v1"
        t.cell(1, 1).text = "v2"
        if mark_first_row_as_header:
            tr = t.rows[0]._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = tr.makeelement(qn('w:trPr'), {})
                tr.insert(0, trPr)
            trPr.append(trPr.makeelement(qn('w:tblHeader'), {}))
        doc.save(path)
        return path

    def test_data_table_first_row_stays_td(self):
        """A table without <w:tblHeader/> must keep its first row as <td>."""
        docx = self._build_table_docx("data.docx", mark_first_row_as_header=False)
        c = WordConverter()
        results, errors = c.convert(docx, "html", self.out_dir)
        self.assertEqual(errors, [])
        content = Path(results[0][1]).read_text(encoding="utf-8")
        import re
        m = re.search(r'<table>.*?</table>', content, re.DOTALL)
        self.assertIsNotNone(m, "no <table> in output")
        table = m.group()
        self.assertNotIn("<th>", table,
                         "BUG 6 regression: data row mis-marked as <th>")
        self.assertIn("<td>", table)

    def test_header_table_first_row_is_th(self):
        """A table with <w:tblHeader/> must render first row as <th>."""
        docx = self._build_table_docx("hdr.docx", mark_first_row_as_header=True)
        c = WordConverter()
        results, errors = c.convert(docx, "html", self.out_dir)
        self.assertEqual(errors, [])
        content = Path(results[0][1]).read_text(encoding="utf-8")
        import re
        m = re.search(r'<table>.*?</table>', content, re.DOTALL)
        self.assertIsNotNone(m, "no <table> in output")
        table = m.group()
        self.assertIn("<th>", table, "header row lost <th>")
        self.assertIn("<thead>", table, "header row lost <thead>")
        self.assertIn("<tbody>", table, "body row lost <tbody>")

    def test_convert_first_row_to_header_removed(self):
        """The buggy helper must no longer be reachable on WordConverter."""
        self.assertFalse(
            hasattr(WordConverter, "_convert_first_row_to_header"),
            "BUG 6 regression: _convert_first_row_to_header still present"
        )


if __name__ == '__main__':
    unittest.main()
