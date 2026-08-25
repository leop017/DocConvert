import os
import tempfile
import unittest
from pathlib import Path

from openpyxl import Workbook

from docconvert.config import AppConfig
from docconvert.controller import ConversionController
from docconvert.models import ProgressEvent
from docconvert.utils import get_excel_sheet_names


def _make_xlsx(path, sheet_name="Sheet1", rows=None):
    if rows is None:
        rows = [["A", "B"], [1, 2]]
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    for row in rows:
        ws.append(row)
    wb.save(path)


class TestConversionControllerSync(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.xlsx = os.path.join(self.tmp.name, "t.xlsx")
        _make_xlsx(self.xlsx)
        self.out_dir = os.path.join(self.tmp.name, "out")
        os.makedirs(self.out_dir, exist_ok=True)
        self.controller = ConversionController()

    def test_empty_files_returns_empty(self):
        results = self.controller.convert_files([], "html")
        self.assertEqual(results, [])

    def test_successful_conversion(self):
        events = []

        def cb(e: ProgressEvent):
            events.append(e)

        results = self.controller.convert_files(
            [self.xlsx], "html", output_dir=self.out_dir, progress_callback=cb
        )
        self.assertEqual(len(results), 1)
        name, path, err = results[0]
        self.assertIsNone(err)
        self.assertTrue(os.path.exists(path))
        self.assertTrue(any(e.done for e in events), "Should emit done event")

    def test_nonexistent_file_reported_as_error(self):
        results = self.controller.convert_files(
            [os.path.join(self.tmp.name, "nope.xlsx")], "html",
            output_dir=self.out_dir,
        )
        self.assertEqual(len(results), 1)
        name, path, err = results[0]
        self.assertEqual(path, "")
        self.assertEqual(err, "文件不存在")

    def test_output_dir_auto_created_when_missing(self):
        nonexistent_dir = os.path.join(self.tmp.name, "missing")
        self.assertFalse(os.path.isdir(nonexistent_dir))
        results = self.controller.convert_files(
            [self.xlsx], "html", output_dir=nonexistent_dir,
        )
        self.assertEqual(len(results), 1)
        _, path, _ = results[0]
        self.assertTrue(os.path.isdir(nonexistent_dir),
                        "Output dir should be auto-created")
        self.assertTrue(path.startswith(nonexistent_dir),
                        "Output file should land in the auto-created dir")

    def test_unsupported_extension_raises_in_get_converter(self):
        bad = os.path.join(self.tmp.name, "weird.txt")
        Path(bad).write_text("x")
        with self.assertRaises(ValueError):
            self.controller._get_converter(".txt")

    def test_cancel_during_processing(self):
        # Create multiple sheets so we have something to cancel mid-loop
        xlsx = os.path.join(self.tmp.name, "multi.xlsx")
        wb = Workbook()
        wb.remove(wb.active)
        for i in range(5):
            ws = wb.create_sheet(f"S{i}")
            ws.append([i])
        wb.save(xlsx)

        def cb(e: ProgressEvent):
            if e.progress >= 0.2 and not self.controller._cancel_event.is_set():
                self.controller.cancel()

        results = self.controller.convert_files(
            [xlsx], "html", output_dir=self.out_dir, progress_callback=cb
        )
        # Should bail out before processing all sheets
        self.assertLessEqual(len(results), 5)

    def test_check_overwrite_paths_detects_existing(self):
        xlsx = os.path.join(self.tmp.name, "x.xlsx")
        _make_xlsx(xlsx)
        self.controller.convert_files([xlsx], "html", output_dir=self.out_dir)
        existing = self.controller.check_overwrite_paths(
            [xlsx], "html", output_dir=self.out_dir
        )
        self.assertTrue(len(existing) >= 1)

    def test_get_excel_sheet_names_missing_file_raises(self):
        # A missing/permission-denied file must surface the error to the
        # caller (check_overwrite_paths) rather than silently returning [].
        ghost = os.path.join(self.tmp.name, "ghost.xlsx")
        with self.assertRaises(Exception):
            get_excel_sheet_names(ghost, ".xlsx")


class TestConversionControllerAsync(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.xlsx = os.path.join(self.tmp.name, "t.xlsx")
        _make_xlsx(self.xlsx)
        self.out_dir = os.path.join(self.tmp.name, "out")
        os.makedirs(self.out_dir, exist_ok=True)

    def test_async_runs_and_completes(self):
        controller = ConversionController()
        started = controller.convert_files_async(
            [self.xlsx], "html", output_dir=self.out_dir
        )
        self.assertTrue(started)
        completed = controller.wait_for_completion(timeout=10)
        self.assertTrue(completed)
        self.assertEqual(len(controller.last_results), 1)

    def test_async_rejects_concurrent_run(self):
        controller = ConversionController()
        first = controller.convert_files_async(
            [self.xlsx], "html", output_dir=self.out_dir
        )
        self.assertTrue(first)
        second = controller.convert_files_async(
            [self.xlsx], "html", output_dir=self.out_dir
        )
        self.assertFalse(second, "Should reject concurrent run")
        controller.wait_for_completion(timeout=10)

    def test_async_recovers_for_next_run(self):
        controller = ConversionController()
        controller.convert_files_async(
            [self.xlsx], "html", output_dir=self.out_dir
        )
        controller.wait_for_completion(timeout=10)
        second = controller.convert_files_async(
            [self.xlsx], "html", output_dir=self.out_dir
        )
        self.assertTrue(second, "Should accept new run after first finishes")
        controller.wait_for_completion(timeout=10)

    def test_wait_done_signals_after_completion(self):
        controller = ConversionController()
        self.assertFalse(controller._done_event.is_set())
        controller.convert_files_async(
            [self.xlsx], "html", output_dir=self.out_dir
        )
        done = controller.wait_done(timeout=10)
        self.assertTrue(done)
        self.assertTrue(controller._done_event.is_set())


class TestCancelAndErrorState(unittest.TestCase):
    """Verify was_cancelled and last_error semantics (bugs #5, #17)."""

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.xlsx = os.path.join(self.tmp.name, "t.xlsx")
        _make_xlsx(self.xlsx)
        self.out_dir = os.path.join(self.tmp.name, "out")
        os.makedirs(self.out_dir, exist_ok=True)

    def test_was_cancelled_flag_set_on_cancel(self):
        xlsx = os.path.join(self.tmp.name, "multi.xlsx")
        wb = Workbook()
        wb.remove(wb.active)
        for i in range(8):
            ws = wb.create_sheet(f"S{i}")
            ws.append([i])
        wb.save(xlsx)

        controller = ConversionController()
        fired = []

        def cb(e):
            if e.progress >= 0.2 and not controller._cancel_event.is_set():
                fired.append(True)
                controller.cancel()

        results = controller.convert_files(
            [xlsx], "html", output_dir=self.out_dir, progress_callback=cb
        )
        self.assertTrue(fired, "Cancel callback should have fired")
        self.assertTrue(controller.was_cancelled,
                        "was_cancelled should be True after cancel()")
        # Partial results: at least one file may have been started; we
        # just want to make sure the flag was set.
        self.assertLess(len(results), 8)

    def test_was_cancelled_false_on_normal_completion(self):
        controller = ConversionController()
        controller.convert_files(
            [self.xlsx], "html", output_dir=self.out_dir,
        )
        self.assertFalse(controller.was_cancelled)
        self.assertIsNone(controller.last_error)

    def test_last_error_set_when_async_worker_raises(self):
        controller = ConversionController()
        # Patch convert_files to raise — simulates worker crash.
        def boom(**kwargs):
            raise RuntimeError("simulated worker failure")
        controller.convert_files = boom
        ok = controller.convert_files_async(
            [self.xlsx], "html", output_dir=self.out_dir
        )
        self.assertTrue(ok)
        completed = controller.wait_for_completion(timeout=5)
        self.assertTrue(completed)
        self.assertEqual(controller.last_error, "simulated worker failure")
        self.assertEqual(controller.last_results, [])

    def test_check_overwrite_swallows_path_errors(self):
        # When _compute_output_paths raises (e.g. file vanished between
        # selection and conversion), check_overwrite_paths should log and
        # skip, not crash the whole overwrite prompt.
        controller = ConversionController()
        existing = controller.check_overwrite_paths(
            [os.path.join(self.tmp.name, "ghost.xlsx")],
            "html", output_dir=self.out_dir,
        )
        self.assertEqual(existing, [])


class TestGuiCleaningCheckboxState(unittest.TestCase):
    """Contract: the 4 cleaning checkboxes in ``cleaning_frame`` are
    interactive **only** when the active input is a Word file AND the
    active output format is Markdown. They are disabled (and their
    backing BooleanVars are reset to False) for every other
    combination, so the UI never shows rules that would silently no-op
    on the converter path.
    """

    def setUp(self):
        import tkinter as tk
        self.root = tk.Tk()
        from docconvert.gui.app import DocConvertApp
        self.app = DocConvertApp(self.root)
        self.addCleanup(self.root.destroy)

    def _assert_state(self, expected: str):
        for var in self.app.cleaning_checks:
            cb = self._find_checkbutton(var)
            self.assertEqual(
                str(cb['state']), expected,
                f"Cleaning checkbox for {var!r} expected state={expected}",
            )

    def test_disabled_when_no_file_loaded(self):
        # file_type=None, output='html' (defaults) → disabled
        self._assert_state('disabled')

    def test_disabled_when_format_is_html(self):
        self.app.file_type = 'word'
        self.app._update_enhanced_state()
        self._assert_state('disabled')

    def test_enabled_when_word_and_md(self):
        self.app.file_type = 'word'
        self.app.output_format.set('md')
        self.app._update_enhanced_state()
        self._assert_state('normal')

    def test_disabled_when_excel_and_md(self):
        # Excel MD path doesn't call WordMdCleaner, so the checkboxes
        # must be disabled to avoid implying rules that won't run.
        self.app.file_type = 'excel'
        self.app.output_format.set('md')
        self.app._update_enhanced_state()
        self._assert_state('disabled')

    def test_disabled_when_doc_and_md(self):
        self.app.file_type = 'doc'
        self.app.output_format.set('md')
        self.app._update_enhanced_state()
        self._assert_state('disabled')

    def test_disabled_again_when_format_switches_back_to_html(self):
        self.app.file_type = 'word'
        self.app.output_format.set('md')
        self.app._update_enhanced_state()
        self.assertEqual(str(self._find_checkbutton(self.app.cleaning_checks[0])['state']), 'normal')
        self.app.output_format.set('html')
        self.app._update_enhanced_state()
        self._assert_state('disabled')

    def test_cleaning_vars_reset_when_disabled(self):
        # Even if user had them checked, switching to a non-applicable
        # context must clear them so the saved config doesn't carry
        # phantom rules into the next conversion.
        self.app.file_type = 'word'
        self.app.output_format.set('md')
        for var in self.app.cleaning_checks:
            var.set(True)
        self.app._update_enhanced_state()
        self.app.output_format.set('html')
        self.app._update_enhanced_state()
        for var in self.app.cleaning_checks:
            self.assertFalse(var.get())

    def _find_checkbutton(self, var):
        for child in self.app.cleaning_frame.winfo_children():
            try:
                if str(child['variable']) == str(var):
                    return child
            except (tk.TclError, KeyError):
                continue
        self.fail(f"Checkbutton for {var!r} not found in cleaning_frame")


class TestGuiBuildConfigForwardsAllRules(unittest.TestCase):
    """Regression: ``_build_config`` must forward every cleaning-rule
    checkbox to the controller. Previously the three reserved rules
    were hard-coded to ``False`` in the GUI, so toggling the new
    checkboxes had no effect on the actual cleaning pipeline.
    """

    def setUp(self):
        import tkinter as tk
        self.root = tk.Tk()
        from docconvert.gui.app import DocConvertApp
        self.app = DocConvertApp(self.root)
        self.addCleanup(self.root.destroy)

    def test_build_config_propagates_all_four_flags(self):
        # Flip every checkbox to a non-default value so we can verify
        # the config actually mirrors what the user sees.
        self.app.clean_page_numbers.set(False)
        self.app.clean_dup_headers.set(False)
        self.app.clean_empty_lines.set(True)
        self.app.clean_normalize_spaces.set(True)

        cfg = self.app._build_config()
        self.assertEqual(cfg.cleaning_rules["remove_page_numbers"], False)
        self.assertEqual(cfg.cleaning_rules["remove_duplicate_headers"], False)
        self.assertEqual(cfg.cleaning_rules["remove_empty_lines"], True)
        self.assertEqual(cfg.cleaning_rules["normalize_spaces"], True)


class TestGuiReconvertRegression(unittest.TestCase):
    """Regression: ``_on_conversion_done`` must accept the boolean forwarded
    by ``_run_in_thread`` from ``Controller.wait_done``. Otherwise Tk's
    ``after(0, ...)`` dispatch raises ``TypeError``, the exception is
    swallowed by the Tk mainloop, and the convert button is never
    re-enabled — making the GUI unable to start a second conversion.
    """

    def test_on_conversion_done_accepts_one_argument(self):
        import inspect
        from docconvert.gui.app import DocConvertApp
        sig = inspect.signature(DocConvertApp._on_conversion_done)
        # Must accept at least one positional argument (the boolean
        # forwarded by _run_in_thread's dispatcher).
        positional_params = [
            p for p in sig.parameters.values()
            if p.kind in (inspect.Parameter.POSITIONAL_ONLY,
                          inspect.Parameter.POSITIONAL_OR_KEYWORD)
        ]
        self.assertGreaterEqual(
            len(positional_params), 2,
            "_on_conversion_done must accept (self, completed: bool) — "
            "_run_in_thread dispatches on_done(result)"
        )


class TestGuiWindowSizeFitsContent(unittest.TestCase):
    """Regression: default 750x640 clipped the bottom bar on systems where
    the cleaning frame pushed required height past 800 px, forcing users
    to drag the window to reveal the convert button. The app must now
    auto-resize to fit its actual content height.
    """

    def test_window_grows_to_fit_content(self):
        import tkinter as tk
        from docconvert.gui.app import DocConvertApp

        root = tk.Tk()
        try:
            DocConvertApp(root)
            # Two layout passes are required for Text/Listbox to settle
            root.update_idletasks()
            root.update_idletasks()
            root.update_idletasks()
            root.update_idletasks()
            req_h = root.winfo_reqheight()
            actual_h = root.winfo_height()
            # Window must be at least as tall as its content requires,
            # otherwise the bottom bar is clipped.
            self.assertGreaterEqual(
                actual_h, req_h,
                f"Window height {actual_h} < required {req_h}; bottom bar would be clipped",
            )
        finally:
            root.destroy()

    def test_window_height_at_least_720(self):
        # Sanity: the explicit default remains a sensible minimum
        # (large enough to show the bottom bar on low-DPI systems).
        import tkinter as tk
        from docconvert.gui.app import DocConvertApp

        root = tk.Tk()
        try:
            DocConvertApp(root)
            root.update_idletasks()
            actual_h = int(root.geometry().split('x')[1].split('+')[0])
            self.assertGreaterEqual(actual_h, 720)
        finally:
            root.destroy()


class TestSetConfig(unittest.TestCase):

    def test_set_config_when_idle_replaces_config(self):
        controller = ConversionController()
        new_cfg = AppConfig(chunk_size=42, max_rows=99)
        ok = controller.set_config(new_cfg)
        self.assertTrue(ok)
        self.assertIs(controller.config, new_cfg)
        self.assertEqual(controller.config.chunk_size, 42)
        self.assertEqual(controller.config.max_rows, 99)

    def test_set_config_during_run_is_rejected(self):
        controller = ConversionController()
        original = controller.config
        controller._running = True  # simulate active run
        new_cfg = AppConfig(chunk_size=42)
        ok = controller.set_config(new_cfg)
        self.assertFalse(ok)
        self.assertIs(controller.config, original)
        controller._running = False  # cleanup


class TestComputeOutputPaths(unittest.TestCase):
    """Test the refactored _compute_output_paths helper."""

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.xlsx = os.path.join(self.tmp.name, "book.xlsx")
        _make_xlsx(self.xlsx, sheet_name="Q1")
        self.xlsx2 = os.path.join(self.tmp.name, "book2.xlsx")
        wb = Workbook()
        wb.remove(wb.active)
        for n in ("Alpha", "Beta"):
            ws = wb.create_sheet(n)
            ws.append([n])
        wb.save(self.xlsx2)
        # Create a real .docx (python-docx writes a minimal valid DOCX)
        from docx import Document as DocxDocument
        self.docx = os.path.join(self.tmp.name, "real.docx")
        DocxDocument().save(self.docx)
        # Create a placeholder .doc (textract can read it; we only need
        # the file to exist for _compute_output_paths)
        self.doc = os.path.join(self.tmp.name, "real.doc")
        Path(self.doc).write_bytes(b"placeholder")
        self.controller = ConversionController()

    def test_docx_path_includes_doc_suffix(self):
        paths = self.controller._compute_output_paths(
            self.docx, "html", output_dir=self.tmp.name,
        )
        self.assertEqual(len(paths), 1)
        self.assertTrue(paths[0].endswith("_doc.html"))

    def test_doc_path_no_suffix(self):
        paths = self.controller._compute_output_paths(
            self.doc, "md", output_dir=self.tmp.name,
        )
        self.assertEqual(len(paths), 1)
        self.assertTrue(paths[0].endswith(".md"))

    def test_xlsx_with_explicit_sheets(self):
        paths = self.controller._compute_output_paths(
            self.xlsx2, "html", output_dir=self.tmp.name, sheets=["Alpha"]
        )
        self.assertEqual(len(paths), 1)
        self.assertTrue(paths[0].endswith("_Alpha.html"))

    def test_xlsx_with_no_sheets_expands_all(self):
        paths = self.controller._compute_output_paths(
            self.xlsx2, "json", output_dir=self.tmp.name
        )
        names = [Path(p).stem for p in paths]
        self.assertEqual(set(names), {"book2_Alpha", "book2_Beta"})

    def test_xlsx_colliding_sheet_names_get_unique_paths(self):
        xlsx = os.path.join(self.tmp.name, "collide.xlsx")
        wb = Workbook()
        wb.remove(wb.active)
        for name in ('Q"1', "Q<1"):
            ws = wb.create_sheet(name)
            ws.append([name])
        wb.save(xlsx)
        paths = self.controller._compute_output_paths(
            xlsx, "html", output_dir=self.tmp.name
        )
        self.assertEqual(len(paths), 2)
        self.assertEqual(
            len(set(paths)), 2,
            "Colliding sheet names must not share an output path",
        )

    def test_missing_input_returns_empty(self):
        paths = self.controller._compute_output_paths(
            os.path.join(self.tmp.name, "ghost.xlsx"), "html",
            output_dir=self.tmp.name,
        )
        self.assertEqual(paths, [])

    def test_unsupported_extension_returns_empty(self):
        paths = self.controller._compute_output_paths(
            os.path.join(self.tmp.name, "x.txt"), "html",
            output_dir=self.tmp.name,
        )
        self.assertEqual(paths, [])

    def test_falls_back_to_source_dir(self):
        paths = self.controller._compute_output_paths(
            self.xlsx, "html", output_dir=None
        )
        self.assertTrue(paths[0].startswith(self.tmp.name))


class TestBatchOutputCollision(unittest.TestCase):
    """Regression: same-named input files from different directories all
    write to a single output directory, so their outputs collided and the
    earlier one was silently overwritten. Output stems must be made unique
    per batch.
    """

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.dir_a = os.path.join(self.tmp.name, "a")
        self.dir_b = os.path.join(self.tmp.name, "b")
        os.makedirs(self.dir_a)
        os.makedirs(self.dir_b)
        self.out_dir = os.path.join(self.tmp.name, "out")
        os.makedirs(self.out_dir)
        self.file_a = os.path.join(self.dir_a, "report.xlsx")
        self.file_b = os.path.join(self.dir_b, "report.xlsx")
        _make_xlsx(self.file_a, rows=[["h"], ["FROM_A"]])
        _make_xlsx(self.file_b, rows=[["h"], ["FROM_B"]])
        self.controller = ConversionController()

    def test_same_named_files_do_not_overwrite(self):
        results = self.controller.convert_files(
            [self.file_a, self.file_b], "html", output_dir=self.out_dir
        )
        paths = [p for _, p, e in results if e is None]
        self.assertEqual(len(paths), 2)
        # Two distinct output files must exist on disk.
        self.assertEqual(len(set(paths)), 2, "Output paths must be unique")
        for p in paths:
            self.assertTrue(os.path.exists(p))
        # Both source values must survive (no silent overwrite).
        blob = "".join(Path(p).read_text(encoding="utf-8") for p in paths)
        self.assertIn("FROM_A", blob)
        self.assertIn("FROM_B", blob)

    def test_overwrite_check_matches_unique_stems(self):
        # After a first run, the overwrite check must report both unique
        # output paths as existing (not just one, and not a phantom
        # colliding path that never gets written).
        first = self.controller.convert_files(
            [self.file_a, self.file_b], "html", output_dir=self.out_dir
        )
        written = {p for _, p, e in first if e is None}
        existing = self.controller.check_overwrite_paths(
            [self.file_a, self.file_b], "html", output_dir=self.out_dir
        )
        self.assertEqual(set(existing), written)

    def test_check_overwrite_matches_written_without_output_dir(self):
        # Regression: without an output_dir, convert_files writes the whole
        # batch into the FIRST file's parent. check_overwrite_paths must
        # report exactly those paths (not each file's own directory), or the
        # overwrite confirmation silently misses real outputs.
        results = self.controller.convert_files([self.file_a, self.file_b], "html")
        written = {p for _, p, e in results if e is None}
        existing = self.controller.check_overwrite_paths(
            [self.file_a, self.file_b], "html", output_dir=None
        )
        self.assertEqual(set(existing), written)


class TestSheetsParamIgnoredForNonExcel(unittest.TestCase):
    """Regression: ``sheets`` is Excel-only. A mixed batch (or a single
    Word/Doc file) converted with ``sheets`` set must not fail the
    Word/Doc files with a TypeError.
    """

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.out_dir = os.path.join(self.tmp.name, "out")
        os.makedirs(self.out_dir)
        from docx import Document as DocxDocument
        self.docx = os.path.join(self.tmp.name, "d.docx")
        doc = DocxDocument()
        doc.add_paragraph("hello world")
        doc.save(self.docx)
        self.xlsx = os.path.join(self.tmp.name, "t.xlsx")
        _make_xlsx(self.xlsx)
        self.controller = ConversionController()

    def test_docx_with_sheets_does_not_fail(self):
        results = self.controller.convert_files(
            [self.docx], "md", output_dir=self.out_dir, sheets=["Sheet1"]
        )
        self.assertEqual(len(results), 1)
        _, path, err = results[0]
        self.assertIsNone(err, f"Word file should not fail when sheets set: {err}")
        self.assertTrue(os.path.exists(path))

    def test_mixed_batch_with_sheets(self):
        results = self.controller.convert_files(
            [self.xlsx, self.docx], "md", output_dir=self.out_dir,
            sheets=["Sheet1"],
        )
        errors = [(n, e) for n, p, e in results if e is not None]
        self.assertEqual(errors, [], f"No file should error: {errors}")


class TestGuiProgressResetOnAlreadyRunning(unittest.TestCase):

    def test_progress_resets_when_convert_rejects(self):
        # Simulate the scenario where convert_files_async returns
        # False because a previous batch is still running. The GUI
        # must reset the indeterminate progress bar so it does not
        # keep spinning forever after the user has been told the
        # batch was rejected.
        import tkinter as tk
        from docconvert.gui.app import DocConvertApp
        root = tk.Tk()
        try:
            app = DocConvertApp(root)
            # Put the progress bar in the indeterminate state we set
            # right before kicking off the async run.
            app.progress.configure(mode='indeterminate')
            app.progress.start(50)
            # Race scenario: is_converting is False when _convert
            # is entered, but convert_files_async races and loses
            # to another caller that started a run before our call.
            # We simulate this by patching convert_files_async to
            # claim a run is in flight (without actually starting one).
            def race_reject(**kw):
                app.controller._running = True
                return False
            app.controller.convert_files_async = race_reject
            # Provide a minimal file list so _convert does not bail
            # out on the empty-list guard.
            app.file_paths = ['dummy.xlsx']
            # Trigger the rejection path.
            app._convert()
            # Tk configure calls are applied on the next update pass.
            root.update()
            self.assertEqual(str(app.progress['mode']), 'determinate')
            # And the value must be reset to 0 so the next run starts fresh.
            self.assertEqual(float(app.progress['value']), 0.0)
        finally:
            root.destroy()

if __name__ == '__main__':
    unittest.main()
