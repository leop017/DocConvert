from __future__ import annotations

import queue
import tkinter as tk
from pathlib import Path
from threading import Thread
from tkinter import ttk, filedialog, messagebox
from typing import Optional

from docx import Document as DocxDocument

from docconvert.config import DEFAULT_CONFIG, AppConfig
from docconvert.controller import ConversionController
from docconvert.logger import get_logger
from docconvert.models import ProgressEvent
from docconvert.utils import decode_text, get_excel_sheet_names


ALL_EXTS = {'.xlsx', '.xls', '.docx', '.doc'}
ALL_EXTS_LIST = sorted(ALL_EXTS)
ALL_EXTS_PATTERN = ' '.join(f'*{e}' for e in ALL_EXTS_LIST)


# ── Modern color palette ──────────────────────────────────────────────
COLORS = {
    'bg':           '#f0f2f5',
    'card_bg':      '#ffffff',
    'accent':       '#2962ff',
    'accent_hover': '#1e4bd8',
    'accent_light': '#e3edff',
    'success':      '#2e7d32',
    'error':        '#c62828',
    'text':         '#1a1a2e',
    'text_sec':     '#5f6368',
    'border':       '#dadce0',
    'preview_bg':   '#fafbfc',
    'listbox_bg':   '#ffffff',
    'listbox_sel':  '#e3edff',
    'title_bar':    '#2962ff',
}


class DocConvertApp:

    def __init__(self, root):
        self.root = root
        self.logger = get_logger()
        self.controller = ConversionController(DEFAULT_CONFIG)

        self.root.title('DocConvert - 文档转换工具')
        self.root.geometry('780x760')
        self.root.resizable(True, True)
        self.root.minsize(680, 600)
        self.root.configure(bg=COLORS['bg'])

        # Set window icon
        self._set_window_icon()

        self.input_file = tk.StringVar()
        self.file_type: Optional[str] = None
        self.file_paths: list[str] = []
        self.sheet_names: list[str] = []
        self.selected_sheet = tk.StringVar()
        self.output_format = tk.StringVar(value='html')
        self.enhanced_md = tk.BooleanVar(value=False)
        self.clean_page_numbers = tk.BooleanVar(value=True)
        self.clean_dup_headers = tk.BooleanVar(value=True)
        self.clean_empty_lines = tk.BooleanVar(value=True)
        self.clean_normalize_spaces = tk.BooleanVar(value=True)
        self.output_dir = tk.StringVar()
        self.convert_all = tk.BooleanVar(value=False)
        self._call_queue: queue.Queue = queue.Queue()
        self._destroying = False
        self._create_styles()
        self._create_widgets()
        self._create_menu()
        self._setup_hover_effects()
        self.root.update_idletasks()
        self.root.update_idletasks()
        req_h = self.root.winfo_reqheight()
        if req_h > 760:
            self.root.geometry(f'780x{req_h + 20}')
        self.root.update_idletasks()
        self.root.update_idletasks()
        self.root.protocol('WM_DELETE_WINDOW', self._on_close)
        self.root.after(50, self._poll_pending)

    def _set_window_icon(self):
        """Set the window icon from the .ico file next to the exe or source."""
        icon_candidates = [
            Path(__file__).parent.parent.parent / 'dist' / 'DocConvert.ico',
            Path(__file__).parent.parent.parent / 'DocConvert.ico',
        ]
        for icon_path in icon_candidates:
            if icon_path.exists():
                try:
                    self.root.iconbitmap(str(icon_path))
                    return
                except tk.TclError:
                    pass
        # Fallback: try to generate a photo icon programmatically
        try:
            self._set_photo_icon()
        except Exception:
            pass

    def _set_photo_icon(self):
        """Create a small blue icon as a PhotoImage fallback."""
        size = 32
        img = tk.PhotoImage(width=size, height=size)
        # Draw a simple blue rounded-square icon
        for y in range(size):
            for x in range(size):
                # Check if inside rounded rectangle
                margin = 2
                r = 4
                ix, iy = x - margin, y - margin
                sx, sy = size - 2 * margin, size - 2 * margin
                in_rect = (margin <= x < size - margin) and (margin <= y < size - margin)
                if in_rect:
                    # Corner rounding check
                    corners = [
                        (r, r), (sx - r, r), (r, sy - r), (sx - r, sy - r)
                    ]
                    in_corner = False
                    for cx, cy in corners:
                        dx, dy = ix - cx, iy - cy
                        if dx < 0 and dy < 0 and dx * dx + dy * dy > r * r:
                            in_corner = True
                            break
                    if not in_corner:
                        img.put('#2962ff', (x, y))
        self._icon_photo = img
        self.root.tk.call('wm', 'iconphoto', self.root._w, img)

    def _create_styles(self):
        style = ttk.Style()
        style.theme_use('clam')

        # ── Frame styles ──
        style.configure('Card.TFrame', background=COLORS['card_bg'], relief='flat')
        style.configure('Main.TFrame', background=COLORS['bg'])
        style.configure('TitleBar.TFrame', background=COLORS['title_bar'])
        style.configure('Bottom.TFrame', background=COLORS['bg'])
        style.configure('Input.TFrame', background=COLORS['card_bg'])

        # ── Label styles ──
        style.configure('AppTitle.TLabel',
                        font=('Microsoft YaHei UI', 20, 'bold'),
                        background=COLORS['title_bar'],
                        foreground='#ffffff')
        style.configure('AppSubtitle.TLabel',
                        font=('Microsoft YaHei UI', 9),
                        background=COLORS['title_bar'],
                        foreground='#b3c6ff')
        style.configure('CardTitle.TLabel',
                        font=('Microsoft YaHei UI', 10, 'bold'),
                        background=COLORS['card_bg'],
                        foreground=COLORS['text'])
        style.configure('Field.TLabel',
                        font=('Microsoft YaHei UI', 9),
                        background=COLORS['card_bg'],
                        foreground=COLORS['text_sec'])
        style.configure('Status.TLabel',
                        font=('Microsoft YaHei UI', 9),
                        background=COLORS['bg'],
                        foreground=COLORS['text_sec'])
        style.configure('StatusSuccess.TLabel',
                        font=('Microsoft YaHei UI', 9, 'bold'),
                        background=COLORS['bg'],
                        foreground=COLORS['success'])
        style.configure('StatusError.TLabel',
                        font=('Microsoft YaHei UI', 9, 'bold'),
                        background=COLORS['bg'],
                        foreground=COLORS['error'])

        # ── Button styles ──
        style.configure('Accent.TButton',
                        font=('Microsoft YaHei UI', 10, 'bold'),
                        background=COLORS['accent'],
                        foreground='#ffffff',
                        borderwidth=0,
                        padding=(20, 8))
        style.map('Accent.TButton',
                  background=[('active', COLORS['accent_hover']),
                              ('disabled', '#b0bec5')])

        style.configure('Secondary.TButton',
                        font=('Microsoft YaHei UI', 9),
                        background=COLORS['card_bg'],
                        foreground=COLORS['text'],
                        borderwidth=1,
                        relief='solid',
                        padding=(12, 5))
        style.map('Secondary.TButton',
                  background=[('active', COLORS['accent_light'])])

        style.configure('Small.TButton',
                        font=('Microsoft YaHei UI', 8),
                        padding=(8, 3))

        # ── Entry styles ──
        style.configure('Modern.TEntry',
                        font=('Microsoft YaHei UI', 9),
                        borderwidth=1,
                        relief='solid',
                        padding=5)

        # ── LabelFrame styles ──
        style.configure('Card.TLabelframe',
                        background=COLORS['card_bg'],
                        foreground=COLORS['text'],
                        borderwidth=1,
                        relief='solid',
                        font=('Microsoft YaHei UI', 10, 'bold'))
        style.configure('Card.TLabelframe.Label',
                        font=('Microsoft YaHei UI', 10, 'bold'),
                        background=COLORS['card_bg'],
                        foreground=COLORS['accent'])

        # ── Checkbutton / Radiobutton ──
        style.configure('Modern.TCheckbutton',
                        font=('Microsoft YaHei UI', 9),
                        background=COLORS['card_bg'],
                        foreground=COLORS['text'])
        style.configure('Modern.TRadiobutton',
                        font=('Microsoft YaHei UI', 9),
                        background=COLORS['card_bg'],
                        foreground=COLORS['text'])

        # ── Progressbar ──
        style.configure('Accent.Horizontal.TProgressbar',
                        troughcolor=COLORS['border'],
                        background=COLORS['accent'],
                        thickness=6)

        # ── Combobox ──
        style.configure('Modern.TCombobox',
                        font=('Microsoft YaHei UI', 9),
                        padding=4)

        # ── Separator ──
        style.configure('Grey.TSeparator', background=COLORS['border'])

    def _setup_hover_effects(self):
        """Add hover color changes to accent buttons."""
        self.convert_btn.bind('<Enter>', lambda e: self.convert_btn.configure(
            style='AccentHover.TButton' if self.convert_btn.cget('state') == 'normal' else 'Accent.TButton'))
        self.convert_btn.bind('<Leave>', lambda e: self.convert_btn.configure(style='Accent.TButton'))

        # Define hover style
        style = ttk.Style()
        style.configure('AccentHover.TButton',
                        font=('Microsoft YaHei UI', 10, 'bold'),
                        background=COLORS['accent_hover'],
                        foreground='#ffffff',
                        borderwidth=0,
                        padding=(20, 8))

    @property
    def is_converting(self) -> bool:
        return self.controller.is_running

    def _on_close(self):
        self._destroying = True
        if self.controller.is_running:
            self.controller.cancel()
        self.root.after(200, self._real_close)

    def _real_close(self):
        if not self.controller.is_running:
            self.root.destroy()
        else:
            self.root.after(100, self._real_close)

    def _create_menu(self):
        menubar = tk.Menu(self.root, bg=COLORS['card_bg'], fg=COLORS['text'],
                          activebackground=COLORS['accent_light'], activeforeground=COLORS['accent'])
        self.root.config(menu=menubar)
        help_menu = tk.Menu(menubar, tearoff=0, bg=COLORS['card_bg'], fg=COLORS['text'],
                            activebackground=COLORS['accent_light'], activeforeground=COLORS['accent'])
        menubar.add_cascade(label='帮助', menu=help_menu)
        help_menu.add_command(label='使用说明', command=self._show_help)
        help_menu.add_separator()
        help_menu.add_command(label='关于', command=self._show_about)

    def _poll_pending(self):
        """Drain the worker -> main callback queue on the Tk main thread.

        Tkinter is not thread-safe: ``root.after`` must only ever be called
        from the main thread. Background threads therefore enqueue callbacks
        here and this poller (scheduled via ``after`` on the main thread)
        runs them.
        """
        if self._destroying:
            return
        while True:
            try:
                fn = self._call_queue.get_nowait()
            except queue.Empty:
                break
            try:
                fn()
            except Exception as e:
                self.logger.warning("UI 回调异常： %s", e)
        self.root.after(50, self._poll_pending)

    def _run_in_thread(self, work, on_done=None):
        def _wrapper():
            try:
                result = work()
            except Exception as e:
                self.logger.warning("Background worker error: %s", e)
                result = None
            if on_done is not None:
                def _dispatch(r=result):
                    on_done(r)
                self._call_queue.put(_dispatch)
        Thread(target=_wrapper, daemon=True).start()

    def _show_help(self):
        help_text = (
            '文档转换工具 使用说明：\n\n'
            '1. 点击"浏览"选择单个文件，或"添加文件"批量选择多个\n'
            '2. 在文件列表中点击可预览各文件内容\n'
            '3. 选择输出格式（HTML/Markdown/JSON）\n'
            '4. 选择输出目录（可选，默认为第一个文件的目录）\n'
            '5. 点击"开始转换"处理列表中的所有文件\n\n'
            '提示：\n'
            '• Excel 支持带合并单元格的表格\n'
            '• Word 转换会自动清除页眉/页脚等非结构化内容\n'
            '• 勾选"增强"可获得更好的 Markdown 效果\n'
        )
        messagebox.showinfo('使用说明', help_text)

    def _show_about(self):
        messagebox.showinfo('关于', 'DocConvert v2.0\n\n文档转换工具\n支持 Excel/Word 格式转 HTML/Markdown/JSON')

    def _create_widgets(self):
        # ── Title bar ──
        title_bar = ttk.Frame(self.root, style='TitleBar.TFrame')
        title_bar.pack(fill=tk.X)
        title_inner = ttk.Frame(title_bar, style='TitleBar.TFrame')
        title_inner.pack(padx=20, pady=(14, 12), anchor=tk.W)
        ttk.Label(title_inner, text='DocConvert', style='AppTitle.TLabel').pack(side=tk.LEFT)
        ttk.Label(title_inner, text='   文档转换工具  v2.0', style='AppSubtitle.TLabel').pack(side=tk.LEFT, padx=(8, 0))

        # ── Main content area ──
        main_frame = ttk.Frame(self.root, style='Main.TFrame', padding=(15, 10, 15, 5))
        main_frame.pack(fill=tk.BOTH, expand=True)

        self._build_input_card(main_frame)
        self._build_file_list_card(main_frame)
        self._build_format_card(main_frame)
        self._build_preview_card(main_frame)
        self._build_bottom_bar(main_frame)

    def _build_input_card(self, parent):
        card = ttk.LabelFrame(parent, text=' 文件选择 ', style='Card.TLabelframe', padding=12)
        card.pack(fill=tk.X, pady=(0, 8))

        ttk.Label(card, text='输入文件:', style='Field.TLabel').grid(row=0, column=0, sticky=tk.W, pady=6, padx=(0, 8))
        path_frame = ttk.Frame(card, style='Input.TFrame')
        path_frame.grid(row=0, column=1, sticky=tk.EW, pady=6)
        self.path_entry = ttk.Entry(path_frame, textvariable=self.input_file,
                                    font=('Microsoft YaHei UI', 9), style='Modern.TEntry')
        self.path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 8))
        ttk.Button(path_frame, text='浏览', style='Secondary.TButton',
                   command=self._browse_file, width=8).pack(side=tk.LEFT)

        ttk.Label(card, text='工作表:', style='Field.TLabel').grid(row=1, column=0, sticky=tk.W, pady=6, padx=(0, 8))
        sheet_frame = ttk.Frame(card, style='Input.TFrame')
        sheet_frame.grid(row=1, column=1, sticky=tk.EW, pady=6)
        self.sheet_combo = ttk.Combobox(sheet_frame, textvariable=self.selected_sheet,
                                        state='readonly', width=35,
                                        font=('Microsoft YaHei UI', 9), style='Modern.TCombobox')
        self.sheet_combo.pack(side=tk.LEFT, fill=tk.X, expand=True)

        ttk.Label(card, text='输出目录:', style='Field.TLabel').grid(row=2, column=0, sticky=tk.W, pady=6, padx=(0, 8))
        outdir_frame = ttk.Frame(card, style='Input.TFrame')
        outdir_frame.grid(row=2, column=1, sticky=tk.EW, pady=6)
        self.outdir_entry = ttk.Entry(outdir_frame, textvariable=self.output_dir,
                                      font=('Microsoft YaHei UI', 9), style='Modern.TEntry')
        self.outdir_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 8))
        ttk.Button(outdir_frame, text='选择', style='Secondary.TButton',
                   command=self._browse_output_dir, width=8).pack(side=tk.LEFT, padx=(0, 4))
        ttk.Button(outdir_frame, text='默认', style='Small.TButton',
                   command=self._use_default_dir, width=5).pack(side=tk.LEFT)

        self.batch_check = ttk.Checkbutton(card, text='转换所有工作表',
                                           variable=self.convert_all, command=self._toggle_batch,
                                           style='Modern.TCheckbutton')
        self.batch_check.grid(row=3, column=1, sticky=tk.W, pady=2)

        card.columnconfigure(1, weight=1)

    def _build_file_list_card(self, parent):
        card = ttk.LabelFrame(parent, text=' 文件列表 ', style='Card.TLabelframe', padding=8)
        card.pack(fill=tk.X, pady=(0, 8))

        btn_row = ttk.Frame(card, style='Input.TFrame')
        btn_row.pack(fill=tk.X, pady=(0, 6))
        ttk.Button(btn_row, text='添加文件', style='Secondary.TButton',
                   command=self._add_files, width=10).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_row, text='移除选中', style='Secondary.TButton',
                   command=self._remove_file, width=10).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_row, text='清空', style='Small.TButton',
                   command=self._clear_files, width=6).pack(side=tk.LEFT, padx=2)

        list_frame = ttk.Frame(card, style='Input.TFrame')
        list_frame.pack(fill=tk.X)
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL)
        self.file_listbox = tk.Listbox(list_frame, height=4, font=('Consolas', 9),
                                       bg=COLORS['listbox_bg'], fg=COLORS['text'],
                                       selectbackground=COLORS['listbox_sel'],
                                       selectforeground=COLORS['accent'],
                                       borderwidth=1, relief='solid',
                                       highlightthickness=1, highlightcolor=COLORS['accent'],
                                       yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.file_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.file_listbox.pack(fill=tk.X, expand=True)
        self.file_listbox.bind('<<ListboxSelect>>', self._on_select_file)

    def _build_format_card(self, parent):
        card = ttk.LabelFrame(parent, text=' 输出格式 ', style='Card.TLabelframe', padding=12)
        card.pack(fill=tk.X, pady=(0, 8))

        radio_row = ttk.Frame(card, style='Input.TFrame')
        radio_row.pack(fill=tk.X, pady=(0, 6))

        # Custom-styled radio buttons
        for text, val in [('HTML', 'html'), ('Markdown', 'md'), ('JSON', 'json')]:
            rb = ttk.Radiobutton(radio_row, text=text, variable=self.output_format,
                                 value=val, command=self._update_enhanced_state,
                                 style='Modern.TRadiobutton')
            rb.pack(side=tk.LEFT, padx=20)

        self.enhanced_check = ttk.Checkbutton(
            card, text='增强 Markdown 输出（更好的格式）',
            variable=self.enhanced_md, state='disabled',
            style='Modern.TCheckbutton'
        )
        self.enhanced_check.pack(anchor=tk.W, pady=(4, 0))

        self.cleaning_frame = ttk.LabelFrame(card, text=' Markdown 清洗 (Word→MD) ',
                                             style='Card.TLabelframe', padding=8)
        self.cleaning_frame.pack(fill=tk.X, pady=(8, 0))

        checks_data = [
            (self.clean_page_numbers, '移除页码 ([1] / 第N页 / Page X 等）', 0, 0),
            (self.clean_dup_headers, '移除重复页眉', 0, 1),
            (self.clean_empty_lines, '移除多余空行', 1, 0),
            (self.clean_normalize_spaces, '合并多余空白', 1, 1),
        ]
        for var, text, r, c in checks_data:
            cb = ttk.Checkbutton(self.cleaning_frame, text=text, variable=var,
                                 state='disabled', style='Modern.TCheckbutton')
            cb.grid(row=r, column=c, sticky=tk.W, padx=6, pady=2)

        self.cleaning_checks = [
            self.clean_page_numbers, self.clean_dup_headers,
            self.clean_empty_lines, self.clean_normalize_spaces,
        ]

    def _build_preview_card(self, parent):
        card = ttk.LabelFrame(parent, text=' 预览 ', style='Card.TLabelframe', padding=8)
        card.pack(fill=tk.BOTH, expand=True, pady=(0, 8))

        self.preview_text = tk.Text(
            card, height=6, font=('Consolas', 9),
            bg=COLORS['preview_bg'], fg=COLORS['text'],
            insertbackground=COLORS['text'],
            selectbackground=COLORS['accent_light'],
            borderwidth=0, relief='flat',
            highlightthickness=1, highlightcolor=COLORS['border'],
            state='disabled', padx=8, pady=6
        )
        self.preview_text.pack(fill=tk.BOTH, expand=True)

    def _build_bottom_bar(self, parent):
        bottom = ttk.Frame(parent, style='Bottom.TFrame')
        bottom.pack(fill=tk.X, pady=(0, 5))

        # Left: status
        self.status_label = ttk.Label(bottom, text='就绪', style='Status.TLabel')
        self.status_label.pack(side=tk.LEFT, padx=(5, 0))

        # Right: progress + convert button
        right_frame = ttk.Frame(bottom, style='Bottom.TFrame')
        right_frame.pack(side=tk.RIGHT)

        self.progress = ttk.Progressbar(right_frame, mode='indeterminate', length=180,
                                        style='Accent.Horizontal.TProgressbar')
        self.progress.pack(side=tk.LEFT, padx=(0, 12), pady=8)

        self.convert_btn = ttk.Button(right_frame, text='开始转换', style='Accent.TButton',
                                      command=self._convert)
        self.convert_btn.pack(side=tk.LEFT, pady=4)

    # ── Existing logic (unchanged) ────────────────────────────────────

    def _update_enhanced_state(self):
        is_md = self.output_format.get() == 'md'
        is_word = self.file_type == 'word'
        self.enhanced_check.configure(state='normal' if is_md else 'disabled')
        if not is_md:
            self.enhanced_md.set(False)
        clean_state = 'normal' if (is_md and is_word) else 'disabled'
        for child in self.cleaning_frame.winfo_children():
            try:
                child.configure(state=clean_state)
            except tk.TclError:
                pass
        if not (is_md and is_word):
            for var in self.cleaning_checks:
                var.set(False)

    def _browse_output_dir(self):
        dirname = filedialog.askdirectory(title='选择输出目录')
        if dirname:
            self.output_dir.set(dirname)

    def _use_default_dir(self):
        self.output_dir.set('')

    def _build_config(self) -> AppConfig:
        from dataclasses import replace
        return replace(
            self.controller.config,
            cleaning_rules={
                "remove_page_numbers": self.clean_page_numbers.get(),
                "remove_duplicate_headers": self.clean_dup_headers.get(),
                "remove_empty_lines": self.clean_empty_lines.get(),
                "normalize_spaces": self.clean_normalize_spaces.get(),
            },
        )

    def _update_preview(self, info_type='info', message=''):
        self.preview_text.configure(state='normal')
        self.preview_text.delete(1.0, tk.END)
        colors = {
            'info': COLORS['text'],
            'success': COLORS['success'],
            'error': COLORS['error'],
            'header': COLORS['accent'],
        }
        self.preview_text.insert(tk.END, message, info_type)
        self.preview_text.tag_config('info', foreground=colors['info'])
        self.preview_text.tag_config('success', foreground=colors['success'])
        self.preview_text.tag_config('error', foreground=colors['error'])
        self.preview_text.tag_config('header', font=('Consolas', 9, 'bold'), foreground=colors['header'])
        self.preview_text.configure(state='disabled')

    def _show_preview_file(self, filepath):
        self.preview_text.configure(state='normal')
        self.preview_text.delete(1.0, tk.END)
        BUF = self.controller.config.preview_chars
        PREVIEW_LINES = self.controller.config.preview_lines
        MAX_READ = 1024 * 1024
        content = None
        read_truncated = False
        for enc in ('utf-8', 'gbk'):
            try:
                with open(filepath, 'r', encoding=enc) as f:
                    content = f.read(MAX_READ + 1)
                if len(content) > MAX_READ:
                    content = content[:MAX_READ]
                    read_truncated = True
                break
            except UnicodeDecodeError:
                continue
            except OSError as e:
                self.preview_text.insert(tk.END, f'无法预览文件： {filepath}\n{e}', 'error')
                self.preview_text.configure(state='disabled')
                return
        if content is None:
            try:
                with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read(MAX_READ + 1)
                if len(content) > MAX_READ:
                    content = content[:MAX_READ]
                    read_truncated = True
            except OSError as e:
                self.preview_text.insert(tk.END, f'无法预览文件： {filepath}\n{e}', 'error')
                self.preview_text.configure(state='disabled')
                return
        try:
            total_chars = len(content)
            if total_chars == 0:
                total_lines = 0
            elif content.endswith('\n'):
                total_lines = content.count('\n')
            else:
                total_lines = content.count('\n') + 1
            display = content[:BUF]
            display_lines = display.split('\n')
            preview_lines = display_lines[:PREVIEW_LINES]
            preview = '\n'.join(preview_lines)
            truncated = read_truncated or total_chars > BUF or len(display_lines) > PREVIEW_LINES
            self.preview_text.insert(tk.END, f'预览： {Path(filepath).name}\n', 'header')
            total_str = f'共 {total_lines} 行， {total_chars} 字符'
            if truncated:
                total_str += ' （截断）'
            self.preview_text.insert(tk.END, total_str + '\n\n', 'info')
            self.preview_text.insert(tk.END, preview)
            if truncated:
                self.preview_text.insert(tk.END, '\n\n... （内容已截断）')
        except (OSError, UnicodeDecodeError):
            self.preview_text.insert(tk.END, f'无法预览文件： {filepath}', 'error')
        self.preview_text.configure(state='disabled')

    def _load_by_ext(self, filepath):
        if not Path(filepath).exists():
            messagebox.showerror('错误', f'文件不存在： {filepath}')
            return False
        ext = Path(filepath).suffix.lower()
        if ext in ('.xlsx', '.xls'):
            self.file_type = 'excel'
            if ext == '.xls':
                try:
                    import xlrd
                except ImportError:
                    messagebox.showerror('错误', '处理 .xls 文件需要安装 xlrd 库\n\n请运行： pip install xlrd')
                    return False
            self.input_file.set(filepath)
            self._load_sheets(filepath)
            self._update_enhanced_state()
            return True
        elif ext == '.docx':
            self.file_type = 'word'
            self.input_file.set(filepath)
            self._load_word(filepath)
            self._update_enhanced_state()
            return True
        elif ext == '.doc':
            try:
                import textract
            except ImportError:
                messagebox.showerror('错误', '处理 .doc 文件需要安装 textract 库\n\n请运行： pip install textract')
                return False
            self.file_type = 'doc'
            self.input_file.set(filepath)
            self.sheet_combo['values'] = []
            self.sheet_combo.set('')
            self.sheet_combo.configure(state='disabled')
            self.sheet_names = []
            self.selected_sheet.set('')
            self._update_preview('info', f'.doc 文件： {Path(filepath).name}\n（正在后台提取文本，请稍候...)')
            self._load_doc_async(filepath)
            self._update_enhanced_state()
            return True
        else:
            messagebox.showerror('错误', '不支持的文件格式')
            return False

    def _browse_file(self):
        filename = filedialog.askopenfilename(
            filetypes=[
                ('支持的文件', ALL_EXTS_PATTERN),
                ('Excel 文件', '*.xlsx *.xls'),
                ('Word 文件', '*.docx *.doc'),
                ('所有文件', '*.*'),
            ],
            title='选择文件'
        )
        if filename:
            if Path(filename).suffix.lower() in ALL_EXTS:
                if filename not in self.file_paths:
                    self.file_paths.append(filename)
                self._refresh_file_list()
            self._load_by_ext(filename)

    def _add_files(self):
        files = filedialog.askopenfilenames(
            filetypes=[
                ('支持的文件', ALL_EXTS_PATTERN),
                ('Excel 文件', '*.xlsx *.xls'),
                ('Word 文件', '*.docx *.doc'),
                ('所有文件', '*.*'),
            ],
            title='添加文件'
        )
        if not files:
            return
        for f in files:
            if Path(f).suffix.lower() in ALL_EXTS and f not in self.file_paths:
                self.file_paths.append(f)
        self._refresh_file_list()
        if self.file_paths:
            self._load_by_ext(self.file_paths[0])

    def _remove_file(self):
        sel = self.file_listbox.curselection()
        if not sel:
            return
        idx = sel[0]
        self.file_paths.pop(idx)
        self._refresh_file_list()
        if self.file_paths:
            new_idx = min(idx, len(self.file_paths) - 1)
            self._suppress_select_event = True
            self.file_listbox.selection_clear(0, tk.END)
            self.file_listbox.selection_set(new_idx)
            self._load_by_ext(self.file_paths[new_idx])
            self.root.after_idle(
                lambda: setattr(self, '_suppress_select_event', False)
            )
        else:
            self.input_file.set('')
            self._update_preview('info', '文件列表为空')

    def _clear_files(self):
        self.file_paths.clear()
        self.file_listbox.delete(0, tk.END)
        self.input_file.set('')
        self._update_preview('info', '文件列表已清空')

    def _refresh_file_list(self):
        self.file_listbox.delete(0, tk.END)
        names = [Path(f).name for f in self.file_paths]
        dupes = {n for n in names if names.count(n) > 1}
        for f in self.file_paths:
            p = Path(f)
            label = f'{p.parent.name}/{p.name}' if p.name in dupes else p.name
            self.file_listbox.insert(tk.END, label)

    def _on_select_file(self, event):
        if getattr(self, '_suppress_select_event', False):
            return
        sel = self.file_listbox.curselection()
        if sel:
            idx = sel[0]
            if idx < len(self.file_paths):
                self._load_by_ext(self.file_paths[idx])

    def _load_sheets(self, filepath):
        self.sheet_combo.configure(state='readonly')
        try:
            ext = Path(filepath).suffix.lower()
            self.sheet_names = get_excel_sheet_names(filepath, ext)
            self.selected_sheet.set('')
            if self.sheet_names:
                self.sheet_combo['values'] = self.sheet_names
                if len(self.sheet_names) == 1:
                    self.selected_sheet.set(self.sheet_names[0])
                else:
                    self.sheet_combo.current(0)
                    self.selected_sheet.set(self.sheet_names[0])
                info = f'已加载： {len(self.sheet_names)} 个工作表\n\n工作表列表:\n'
                for i, name in enumerate(self.sheet_names[:10]):
                    info += f'  {i + 1}. {name}\n'
                if len(self.sheet_names) > 10:
                    info += f'  ... 还有 {len(self.sheet_names) - 10} 个\n'
                self._update_preview('success', info)
        except Exception as e:
            messagebox.showerror('错误', f'加载文件失败:\n{str(e)}')
            self.sheet_names = []
            self.sheet_combo['values'] = []
            self.sheet_combo.set('')
            self._update_preview('error', f'加载失败： {str(e)}')
        finally:
            if self.convert_all.get():
                self.sheet_combo.configure(state='disabled')

    def _load_word(self, filepath):
        try:
            doc = DocxDocument(filepath)
            paras = [p for p in doc.paragraphs if p.text.strip()]
            tables = doc.tables
            sections = doc.sections
            self.sheet_names = []
            self.selected_sheet.set('')
            self.sheet_combo['values'] = []
            self.sheet_combo.set('')
            self.sheet_combo.configure(state='disabled')
            info = f'Word 文档： {Path(filepath).name}\n'
            info += f'段落数： {len(paras)}\n'
            info += f'表格数： {len(tables)}\n'
            info += f'节数： {len(sections)}\n\n'
            if sections:
                hf_count = 0
                for sec in sections:
                    if sec.header and any(p.text.strip() for p in sec.header.paragraphs):
                        hf_count += 1
                    if sec.footer and any(p.text.strip() for p in sec.footer.paragraphs):
                        hf_count += 1
                if hf_count:
                    info += '含页眉/页脚 （将被清洗）\n'
            if paras:
                info += '\n--- 预览前10段 ---\n'
                for p in paras[:10]:
                    info += p.text[:120] + '\n'
            self._update_preview('success', info)
        except Exception as e:
            messagebox.showerror('错误', f'加载文档失败:\n{str(e)}')
            self._update_preview('error', f'加载失败： {str(e)}')

    def _load_doc_async(self, filepath: str):
        def _work():
            try:
                import textract
                raw = textract.process(filepath)
                text = decode_text(raw)
                lines = text.split('\n')
                preview_lines = [line for line in lines if line.strip()][:20]
                return (
                    f'.doc 文件： {Path(filepath).name}\n'
                    f'共约 {len(lines)} 行\n\n--- 预览 ---\n'
                    + '\n'.join(preview_lines)
                )
            except Exception:
                return f'.doc 文件： {Path(filepath).name}\n（后台文本提取失败）'

        def _on_done(info):
            if self.file_type == 'doc' and self.input_file.get() == filepath:
                self._update_preview('info', info)

        self._run_in_thread(_work, _on_done)

    def _toggle_batch(self):
        if self.convert_all.get():
            self.sheet_combo.configure(state='disabled')
        else:
            self.sheet_combo.configure(state='readonly')

    def _on_progress(self, event: ProgressEvent):
        self._call_queue.put(lambda e=event: self._handle_progress(e))

    def _handle_progress(self, event: ProgressEvent):
        if event.error:
            self.status_label.config(text=event.error, style='StatusError.TLabel')
            self._update_preview('error', event.error)
        elif event.message:
            self.status_label.config(text=event.message, style='Status.TLabel')

        if event.progress > 0:
            # First determinate update: stop the indeterminate animation
            # before switching modes so the bar doesn't keep cycling.
            if str(self.progress['mode']) == 'indeterminate':
                self.progress.stop()
            self.progress['mode'] = 'determinate'
            self.progress['value'] = int(event.progress * 100)

    def _convert(self):
        if self.is_converting:
            return

        output_fmt = self.output_format.get()

        files_to_process = list(self.file_paths)
        if not files_to_process:
            p = self.input_file.get().strip()
            if p:
                files_to_process = [p]

        if not files_to_process:
            messagebox.showerror('错误', '请选择文件')
            return

        out_dir = self.output_dir.get().strip()

        sheets = None
        if not self.convert_all.get():
            sel = self.selected_sheet.get()
            if sel:
                sheets = [sel]

        self.controller.set_config(self._build_config())

        existing = self.controller.check_overwrite_paths(
            files=files_to_process,
            output_fmt=output_fmt,
            output_dir=out_dir if out_dir else None,
            sheets=sheets,
        )
        if existing:
            preview = '\n'.join(f'  \u2022 {Path(p).name}' for p in existing[:10])
            if len(existing) > 10:
                preview += f'\n  ... 还有 {len(existing) - 10} 个'
            proceed = messagebox.askyesno(
                '覆盖确认',
                f'以下 {len(existing)} 个输出文件已存在，将被覆盖:\n{preview}\n\n是否继续？',
            )
            if not proceed:
                return

        self.convert_btn.configure(state='disabled')
        self.status_label.config(text='加载文件中...')
        self.progress.configure(mode='indeterminate', value=0)
        # Animate the bar during the initial loading phase, before the
        # first determinate progress event arrives (which stops it).
        self.progress.start(12)
        self.root.update()

        started = self.controller.convert_files_async(
            files=files_to_process,
            output_fmt=output_fmt,
            output_dir=out_dir if out_dir else None,
            enhanced_md=self.enhanced_md.get(),
            sheets=sheets,
            progress_callback=self._on_progress,
        )

        if not started:
            self.convert_btn.configure(state='normal')
            self.progress.stop()
            self.progress.configure(mode='determinate', value=0)
            self.status_label.config(text='任务已在运行中', style='StatusError.TLabel')
            return

        self._wait_for_done()

    def _wait_for_done(self):
        self._run_in_thread(self.controller.wait_done, self._on_conversion_done)

    def _on_conversion_done(self, _completed: bool = True):
        self.convert_btn.configure(state='normal')
        # Stop any residual indeterminate animation (e.g. a batch that
        # finished before emitting a determinate progress event).
        if str(self.progress['mode']) == 'indeterminate':
            self.progress.stop()
        self.progress['mode'] = 'determinate'

        if self.controller.last_error:
            err = self.controller.last_error
            self.status_label.config(text='转换失败', style='StatusError.TLabel')
            self._update_preview('error', f'转换失败： {err}')
            messagebox.showerror('错误', f'转换失败:\n{err}')
            return

        if not self.controller.was_cancelled:
            self.progress['value'] = 100
        self._show_conversion_results()

    def _show_conversion_results(self):
        results = self.controller.last_results
        all_results = [(n, p) for n, p, e in results if e is None]
        all_errors = [(n, e) for n, p, e in results if e is not None]
        was_cancelled = self.controller.was_cancelled

        if not results:
            self.status_label.config(text='已取消', style='Status.TLabel')
            self._update_preview('info', '转换已取消')
            return

        total = len(results)
        is_multi = total > 1

        if is_multi:
            msg_lines = [f'处理完成： {len(all_results)}/{total} 个输出']
            if all_results:
                msg_lines.append('')
                msg_lines.append('成功:')
                for n, p in all_results[:20]:
                    msg_lines.append(f'  \u2022 {n}')
                if len(all_results) > 20:
                    msg_lines.append(f'  ... 还有 {len(all_results) - 20} 个')
            if all_errors:
                msg_lines.append('')
                msg_lines.append(f'失败 ({len(all_errors)}):')
                for fname, err in all_errors[:10]:
                    msg_lines.append(f'  \u2022 {fname}: {err}')
                if len(all_errors) > 10:
                    msg_lines.append(f'  ... 还有 {len(all_errors) - 10} 个')

            if was_cancelled:
                status_text = f'已取消 ({len(all_results)}/{total})'
                status_style = 'Status.TLabel'
            elif all_results and not all_errors:
                status_text = f'转换完成 ({len(all_results)}/{total})'
                status_style = 'StatusSuccess.TLabel'
            elif all_results and all_errors:
                status_text = f'部分成功 ({len(all_results)}/{total})'
                status_style = 'Status.TLabel'
            else:
                status_text = f'转换失败 (0/{total})'
                status_style = 'StatusError.TLabel'

            self.status_label.config(text=status_text, style=status_style)
            if all_results:
                self._show_preview_file(all_results[0][1])
            messagebox.showinfo('转换结果', '\n'.join(msg_lines))
        elif all_results:
            name, path = all_results[0]
            self.status_label.config(text='转换完成', style='StatusSuccess.TLabel')
            self._show_preview_file(path)
            ext_map = {'html': 'HTML', 'md': 'Markdown', 'json': 'JSON'}
            messagebox.showinfo(
                '成功',
                f'{ext_map.get(self.output_format.get(), self.output_format.get())} 文件已生成:\n{path}'
            )
        else:
            err_msg = all_errors[0][1] if all_errors else '未知错误'
            self.status_label.config(text='转换失败', style='StatusError.TLabel')
            self._update_preview('error', f'转换失败： {err_msg}')
            messagebox.showerror('错误', f'转换失败:\n{err_msg}')
